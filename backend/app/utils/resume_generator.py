import groq
import os
import time
import logging
from typing import Dict, Any, List, Tuple, Optional
import json
from pathlib import Path
from ..latex.processor import LatexProcessor
import tiktoken
import subprocess
from .. import models
from ..database import SessionLocal, get_job_title_from_cache, save_generation_status
from .resume_assessment_agents import (
                content_quality_agent,
                # formatting_agent,
                skills_agent,
                experience_agent,
                resume_constructor_agent,
                content_quality_task,
                # formatting_task,
                skills_task,
                experience_task,
                resume_construction_task,
                calculate_total_tokens,
                create_llm
            )


logger = logging.getLogger(__name__)

class TokenTracker:
    """Tracks token usage and calculates costs for Groq API calls"""
    
    COST_PER_MILLION_TOKENS = 40  # $20 per million tokens
    
    def __init__(self):
        self.total_input_tokens = 0
        self.total_output_tokens = 0
        self.agent_input_tokens = 0
        self.agent_output_tokens = 0
        self.session_history = []
        self.agent_history = []
        self.encoder = tiktoken.get_encoding("cl100k_base")  # Using OpenAI's tokenizer as approximation
        
    def count_tokens(self, text: str) -> int:
        """Count tokens in a string using the tokenizer"""
        return len(self.encoder.encode(text))
    
    def add_api_call(self, prompt: str, response: Optional[str]) -> Dict[str, Any]:
        """Record a new API call with input and output tokens"""
        input_tokens = self.count_tokens(prompt)
        output_tokens = self.count_tokens(response) if response else 0
        
        self.total_input_tokens += input_tokens * 2
        self.total_output_tokens += output_tokens * 2
        
        call_stats = {
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'input_cost': self.calculate_cost(input_tokens),
            'output_cost': self.calculate_cost(output_tokens),
            'total_cost': self.calculate_cost(input_tokens + output_tokens)
        }
        
        self.session_history.append(call_stats)
        return call_stats
    
    def calculate_cost(self, num_tokens: float) -> float:
        """Calculate cost in dollars for a given number of tokens"""
        return (num_tokens / 1_000_000) * self.COST_PER_MILLION_TOKENS * 120
    
    def add_agent_call(self, agent_name: str, context: str, response: str) -> Dict[str, Any]:
        """Record a new agent call with input and output tokens"""
        input_tokens = self.count_tokens(context)
        output_tokens = self.count_tokens(response)
        
        self.agent_input_tokens += input_tokens
        self.agent_output_tokens += output_tokens
        
        call_stats = {
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'agent': agent_name,
            'input_tokens': input_tokens,
            'output_tokens': output_tokens,
            'input_cost': self.calculate_cost(input_tokens),
            'output_cost': self.calculate_cost(output_tokens),
            'total_cost': self.calculate_cost(input_tokens + output_tokens)
        }
        
        self.agent_history.append(call_stats)
        return call_stats

    def get_total_usage(self) -> Dict[str, Any]:
        """Get total token usage and costs for all API calls and agent calls"""
        total_input_tokens = 1.5* (self.total_input_tokens + self.agent_input_tokens)
        total_output_tokens =1.5* (self.total_output_tokens + self.agent_output_tokens)
        total_tokens = total_input_tokens + total_output_tokens
        
        return {
            'total_input_tokens': self.total_input_tokens,
            'total_output_tokens': self.total_output_tokens * 2,
            'agent_input_tokens': self.agent_input_tokens * 2,
            'agent_output_tokens': self.agent_output_tokens * 3,
            'total_tokens': total_tokens,
            'total_input_cost': self.calculate_cost(total_input_tokens),
            'total_output_cost': self.calculate_cost(total_output_tokens),
            'total_cost': self.calculate_cost(total_tokens),
            'call_history': self.session_history,
            'agent_history': self.agent_history
        }

class ResumeGenerator:
    def __init__(self):
        self.client = groq.Groq(api_key=os.getenv("GROQ_API_KEY"))
        self.latex_processor = LatexProcessor()
        self.token_tracker = TokenTracker()
        # Validate Groq API on initialization
        try:
            model = os.getenv("GROQ_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct")  # Default to mixtral if not set
            if not model:
                raise ValueError("GROQ_MODEL environment variable must be set")

            self.client.chat.completions.create(
                messages=[{"role": "user", "content": "test"}],
                model=model,
                max_tokens=1
            )
            logger.info("Successfully validated Groq API connection")
        except Exception as e:
            logger.error(f"Failed to validate Groq API: {str(e)}")
            raise

    def format_experiences(self, experiences: List[str]) -> str:
        """Format experiences into a readable string."""
        if not experiences:
            return "No previous experience provided"
        return "\n".join([f"- {exp}" for exp in experiences])

    def format_education(self, education: List[str]) -> str:
        """Format education into a readable string."""
        if not education:
            return "No education information provided"
        return "\n".join([f"- {edu}" for edu in education])

    def format_skills(self, skills: List[str]) -> str:
        """Format skills into a readable string."""
        if not skills:
            return "No skills provided"
        return ", ".join(skills)

    def generate_optimized_resume(self, professional_info: Dict[str, Any], job_description: str, skills: Optional[List[str]] = None) -> Tuple[Optional[str], Dict[str, Any]]:
        """
        Generate an optimized resume content using Groq AI and track token usage.
        Returns both the generated content and token usage statistics.
        """
        try:
            # Log input data structure
            logger.info(f"Input data structure for optimization: {json.dumps(professional_info, indent=2)}")

            # Format the professional info for the prompt
            experience_text = self.format_experiences(professional_info.get('past_experiences', []))
            education_text = self.format_education(professional_info.get('education', []))
            skills_text = self.format_skills(professional_info.get('skills', []))
            summary_text = professional_info.get('professional_summary', 'Not provided')
            certifications_text = "\n".join([f"- {cert}" for cert in professional_info.get('certifications', [])])

            # Construct the prompt
            prompt = f"""
            As an expert resume writer, create an optimized resume for the job description provided.
            If no previous experience is provided, create compelling fictional experience that would be ideal for this role.
            Focus on relevant experience and skills that match the job requirements.
            Consider the following skills if provided: {skills}.
            
            Note: When no previous experience exists, generate 2-3 relevant positions with achievements that demonstrate
            the key skills and qualifications needed for the target role.

            Original Professional Information:

            Current Summary:
            {summary_text}

            Experience:
            {experience_text}

            Education:
            {education_text}

            Skills:
            {skills_text}

            Certifications:
            {certifications_text}

            Job Description:
            {job_description}

            Generate an optimized resume with the following sections.
            Use the exact format shown below, including the section headers and markers:

            # Professional Summary
            ===
            Write 2-3 compelling sentences highlighting most relevant qualifications.
            ===

            # Key Skills
            ===
            List 4-6 most relevant skills, each on a new line starting with •
            Example:
            • Skill 1
            • Skill 2
            ===

            # Professional Experience
            ===
            For each position, format as:
            [Title] at [Company], [Duration]
            • Achievement 1
            • Achievement 2
            • Achievement 3
            • Achievement 4

            Example:
            Senior Software Engineer at XYZ Corp, 2020-Present
            • Led development of microservices architecture
            • Implemented automated testing pipeline
            ===

            # Education
            ===
            For each education entry, format as:
            [Degree]
            [Institution]
            [Year]

            Example:
            Bachelor of Science in Computer Science
            University of California, Berkeley
            2014
            ===

            # Certifications
            ===
            List relevant certifications, each on a new line starting with •
            Example:
            • AWS Certified Solutions Architect
            • Professional Scrum Master I
            ===

            # Achievements
            ===
            List standalone achievements not tied to specific jobs, each on a new line starting with •
            Example:
            • Won 1st place in national coding competition
            • Published research paper in IEEE journal
            ===

            # Projects
            ===
            List relevant projects, each on a new line starting with •
            Format each project as:
            • [Project Title] - [Brief Description]
            Technologies: [List of technologies used]

            Example:
            • E-commerce Platform - Built a scalable online marketplace
            Technologies: React, Node.js, MongoDB
            ===

            # Others
            ===
            List any other relevant information, each on a new line starting with •
            Example:
            • Open source contributions
            • Professional memberships
            • Awards and recognitions
            ===

            Important:
            1. Keep the exact section headers (# Section Name)
            2. Keep the === markers before and after each section's content
            4. Make the content highly relevant to the job description
            5. Follow the exact formatting shown in the examples
            6. If a section has no content, include the section with 'None' between the === markers
            """

            # Call Groq API
            model = os.getenv("GROQ_MODEL", "meta-llama/llama-4-scout-17b-16e-instruct")  # Default to mixtral if not set
            if not model:
                raise ValueError("GROQ_MODEL environment variable must be set")

            chat_completion = self.client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert resume writer who specializes in creating targeted resumes that align with specific job descriptions. When no previous experience is provided, create compelling fictional experience that would be ideal for the job description. Always format your response exactly as requested, maintaining section headers and markers."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                model=model,
                temperature=0.7,
                max_tokens=2000,
                top_p=1,
                stream=False
            )

            # Extract the generated content and escape special LaTeX characters
            generated_content = chat_completion.choices[0].message.content
            if generated_content:
                generated_content = generated_content.replace('&', '\&').replace('%', '\%')
                generated_content = generated_content.replace('R&D', 'R\&D')

            # Track token usage
            usage_stats = self.token_tracker.add_api_call(prompt, generated_content)

            # logger.info(f"Optimized resume generated by AI:\n{generated_content or 'No content generated'}")
            logger.info(f"Token usage for this generation: {json.dumps(usage_stats, indent=2)}")

            return generated_content, usage_stats

        except Exception as e:
            logger.error(f"Error generating optimized resume: {str(e)}")
        raise

    def get_existing_resume(self, user_id: int) -> Optional[str]:
        """
        Compile information from all available profile sections and use that data for optimized resume generation.
        Returns the compiled profile data as JSON if found and user preference allows, None otherwise.
        """
        from .. import models
        
        try:
            # First check if user wants to use resume sections as reference
            db = SessionLocal()
            try:
                profile = db.query(models.Profile).filter(models.Profile.user_id == user_id).first()
                if not profile:
                    logger.info(f"No profile found for user {user_id}")
                    return None
                
                # Check user preferences - must be True to use profile sections
                if not profile.use_resume_sections:
                    logger.info(f"User {user_id} has disabled using resume sections as reference")
                    return None
                
                # Compile data from all profile sections
                compiled_data = {
                    "work_experience": [],
                    "education": [],
                    "skills": [],
                    "projects": [],
                    "publications": [],
                    "volunteer_work": [],
                    "summary": profile.summary,
                    "professional_title": profile.professional_title
                }
                
                # Work Experience
                for exp in profile.work_experiences:
                    work_exp = {
                        "position": exp.position,
                        "company": exp.company,
                        "location": exp.location,
                        "start_date": exp.start_date.isoformat() if exp.start_date else None,
                        "end_date": exp.end_date.isoformat() if exp.end_date else None,
                        "current_job": exp.current_job,
                        "description": exp.description,
                        "achievements": exp.achievements or [],
                        "technologies": exp.technologies or []
                    }
                    compiled_data["work_experience"].append(work_exp)
                
                # Education
                for edu in profile.educations:
                    education = {
                        "institution": edu.institution,
                        "degree": edu.degree,
                        "field_of_study": edu.field_of_study,
                        "location": edu.location,
                        "start_date": edu.start_date.isoformat() if edu.start_date else None,
                        "end_date": edu.end_date.isoformat() if edu.end_date else None,
                        "gpa": edu.gpa,
                        "description": edu.description,
                        "achievements": edu.achievements or []
                    }
                    compiled_data["education"].append(education)
                
                # Skills
                for skill in profile.skills:
                    skill_data = {
                        "name": skill.name,
                        "category": skill.category,
                        "proficiency": skill.proficiency,
                        "years_experience": skill.years_experience
                    }
                    compiled_data["skills"].append(skill_data)
                
                # Projects
                for project in profile.projects:
                    project_data = {
                        "name": project.name,
                        "description": project.description,
                        "url": project.url,
                        "github_url": project.github_url,
                        "start_date": project.start_date.isoformat() if project.start_date else None,
                        "end_date": project.end_date.isoformat() if project.end_date else None,
                        "technologies": project.technologies or [],
                        "achievements": project.achievements or []
                    }
                    compiled_data["projects"].append(project_data)
                
                # Publications
                for pub in profile.publications:
                    publication = {
                        "title": pub.title,
                        "publisher": pub.publisher,
                        "publication_date": pub.publication_date.isoformat() if pub.publication_date else None,
                        "url": pub.url,
                        "description": pub.description,
                        "authors": pub.authors or []
                    }
                    compiled_data["publications"].append(publication)
                
                # Volunteer Work
                for vol in profile.volunteer_works:
                    volunteer = {
                        "organization": vol.organization,
                        "role": vol.role,
                        "cause": vol.cause,
                        "location": vol.location,
                        "start_date": vol.start_date.isoformat() if vol.start_date else None,
                        "end_date": vol.end_date.isoformat() if vol.end_date else None,
                        "current_role": vol.current_role,
                        "description": vol.description,
                        "achievements": vol.achievements or []
                    }
                    compiled_data["volunteer_work"].append(volunteer)
                
                # Check if we have any meaningful data
                has_data = any([
                    compiled_data["work_experience"],
                    compiled_data["education"],
                    compiled_data["skills"],
                    compiled_data["projects"],
                    compiled_data["publications"],
                    compiled_data["volunteer_work"],
                    compiled_data["summary"],
                    compiled_data["professional_title"]
                ])
                
                if not has_data:
                    logger.info(f"No profile section data found for user {user_id}")
                    return None
                
                logger.info(f"Using compiled profile sections as reference for user {user_id}")
                logger.info(f"Compiled data includes: {len(compiled_data['work_experience'])} work experiences, "
                          f"{len(compiled_data['education'])} education entries, {len(compiled_data['skills'])} skills, "
                          f"{len(compiled_data['projects'])} projects, {len(compiled_data['publications'])} publications, "
                          f"{len(compiled_data['volunteer_work'])} volunteer experiences")
                
                return json.dumps(compiled_data)
                    
            finally:
                db.close()
            
        except Exception as e:
            logger.error(f"Error compiling profile section data: {str(e)}")
            return None

    def estimate_generation_time(self, job_description: str, has_existing_resume: bool = False) -> int:
        """
        Estimate total generation time based on job description complexity and existing data.
        
        Args:
            job_description: The job description text
            has_existing_resume: Whether user has an existing resume
            
        Returns:
            int: Estimated time in seconds
        """
        base_time = 60 if has_existing_resume else 30  # Base time in seconds
        
        # Add time based on job description complexity
        job_desc_length = len(job_description)
        complexity_factor = min(job_desc_length / 1000 * 30, 60)  # Max 60s additional
        
        # Agent processing time (4 agents * average time per agent)
        agent_time = 4 * 15  # 45 seconds per agent
        
        total_estimate = int(base_time + complexity_factor + agent_time)
        return min(total_estimate, 300)  # Cap at 10 minutes

    async def optimize_resume(self, resume_gen_id:str, professional_info: Dict[str, Any], job_description: str, skills: Optional[List[str]] = None, user_id: Optional[int] = None) -> Dict[str, Any]:
        """
        Main function to optimize the entire resume using assessment agents with progress tracking.
        Returns the optimized content along with token usage statistics.
        """
        try:
            # Initialize progress tracking
            save_generation_status(resume_gen_id, "parsing", 5, "Analyzing job description and requirements...")
            
            # Check for existing resume first
            initial_content = None
            usage_stats = {}
            has_existing_resume = False
            
            if user_id is not None:
                existing_resume = self.get_existing_resume(user_id)
                if existing_resume is not None:
                    initial_content = existing_resume
                    has_existing_resume = True
                    logger.info(f"Used Existing Resume")
            
            # Estimate total time
            estimated_time = self.estimate_generation_time(job_description, has_existing_resume)
            save_generation_status(resume_gen_id, "parsing", 10, "Processing job requirements...", estimated_time - 5)
            
            # Generate initial content using Groq AI if no existing resume
            if initial_content is None:
                save_generation_status(resume_gen_id, "parsing", 15, "Generating initial resume content...", estimated_time - 10)
                initial_content, usage_stats = self.generate_optimized_resume(professional_info, job_description, skills)
                logger.info(f"Used Fake Resume from Groq")
            
            context = f"job description:\n{job_description}\n############\ninitial_content:\n{initial_content}"
            
            # Execute tasks with timeout handling and progress tracking
            def execute_with_timeout(agent, task, timeout=300):
                import concurrent.futures
                max_retries = 3
                base_delay = 1
                
                for attempt in range(max_retries):
                    try:
                        with concurrent.futures.ThreadPoolExecutor() as executor:
                            future = executor.submit(agent.execute_task, task, context=context)
                            try:
                                return future.result(timeout=timeout)
                            except concurrent.futures.TimeoutError:
                                logger.error(f"Timeout executing {agent.role}")
                                raise Exception(f"Timeout error for {agent.role}")
                                
                    except Exception as e:
                        error_str = str(e)
                        if "503" in error_str and "VertexAI" in error_str:
                            logger.warning(f"VertexAI model overloaded (attempt {attempt + 1}/{max_retries})")
                        elif "Timeout" in error_str:
                            logger.error(f"Timeout executing {agent.role}")
                        else:
                            logger.error(f"Error executing {agent.role}: {error_str}")
                            
                        if attempt == max_retries - 1:
                            logger.error(f"Failed after {max_retries} attempts for {agent.role}")
                            # Force fallback to Groq
                            agent.llm = create_llm(model="AGENT")
                            return agent.execute_task(task, context=context)
                            
                        # Exponential backoff
                        delay = base_delay * (2 ** attempt)
                        time.sleep(min(delay, 10))

            # Content Quality Analysis (25-45%)
            save_generation_status(resume_gen_id, "analyzing", 25, "Analyzing content quality and relevance...", estimated_time - 20)
            content_quality_result = execute_with_timeout(content_quality_agent, content_quality_task)
            self.token_tracker.add_agent_call(
                "content_quality_agent",
                context,
                content_quality_result
            )

            # Skills Analysis (45-65%)
            save_generation_status(resume_gen_id, "optimizing", 45, "Optimizing skills alignment with job requirements...", estimated_time - 50)
            skills_result = execute_with_timeout(skills_agent, skills_task)
            self.token_tracker.add_agent_call(
                "skills_agent",
                context,
                skills_result
            )

            # Experience Analysis (65-85%)
            save_generation_status(resume_gen_id, "optimizing", 65, "Enhancing experience descriptions and achievements...", estimated_time - 65)
            experience_result = execute_with_timeout(experience_agent, experience_task)
            self.token_tracker.add_agent_call(
                "experience_agent",
                context,
                experience_result
            )

            # Combine agent outputs
            agent_outputs = f"""
            {content_quality_result}
            \n
            {skills_result}
            \n
            {experience_result}
            """

            # Final Resume Construction (85-95%)
            save_generation_status(resume_gen_id, "constructing", 85, "Constructing final optimized resume...", estimated_time - 80)
            final_resume = execute_with_timeout(resume_constructor_agent, resume_construction_task, timeout=90)

            # Save to database (95-100%)
            save_generation_status(resume_gen_id, "constructing", 95, "Finalizing and saving resume...", 10)
            
            # Save the resume to the database
            if user_id:
                db = SessionLocal()
                try:
                    profile = db.query(models.Profile).filter(models.Profile.user_id == user_id).first()
                    if not profile:
                        raise ValueError(f"No profile found for user_id: {user_id}")
                    
                    # Check existing resume count and delete oldest if needed
                    existing_resumes = db.query(models.Resume)\
                        .filter(models.Resume.profile_id == profile.id)\
                        .order_by(models.Resume.created_at)\
                        .all()
                    
                    if len(existing_resumes) >= 10:
                        oldest_resume = existing_resumes[0]
                        db.delete(oldest_resume)
                        logger.info(f"Deleted oldest resume {oldest_resume.id} to maintain 10-resume limit")
                    
                    # Save new resume
                    db_resume = models.Resume(
                        id = resume_gen_id,
                        profile_id=profile.id,
                        content=final_resume,
                        job_description=job_description,
                        name=f"Resume for id : {resume_gen_id[-4:]}|{get_job_title_from_cache(resume_gen_id)}",
                        version="1.0",
                        status="completed"
                    )
                    db.add(db_resume)
                    db.commit()
                    logger.info(f"Successfully saved resume to database for user {user_id}")
                except Exception as e:
                    logger.error(f"Error saving resume to database: {str(e)}")
                finally:
                    db.close()

            # Format the analysis summary
            from .resume_assessment_agents import format_analysis_summary
            analysis_summary = format_analysis_summary(agent_outputs)

            # Prepare the result data
            result_data = {
                'job_id': resume_gen_id,
                'job_title': get_job_title_from_cache(resume_gen_id) or 'Resume',
                'content': final_resume,  # Frontend expects 'content', not 'ai_content'
                'ai_content': final_resume,  # Keep both for backward compatibility
                'professional_info': professional_info,
                'token_usage': usage_stats,
                'total_usage': self.token_tracker.get_total_usage(),
                'agent_outputs': agent_outputs,
                'analysis_summary': analysis_summary,
                'message': 'Resume generated successfully'
            }

            # Save the result to Redis cache for retrieval FIRST
            from ..database import save_generation_result
            save_generation_result(resume_gen_id, result_data, expiration=3600)
            logger.info(f"Saved generation result to cache for job {resume_gen_id}")

            # ONLY mark as completed AFTER the result is saved
            save_generation_status(resume_gen_id, "completed", 100, "Resume generation completed successfully!", 0)

            return result_data

        except Exception as e:
            logger.error(f"Error optimizing resume: {str(e)}")
            save_generation_status(resume_gen_id, "failed", 0, f"Generation failed: {str(e)}", 0)
            raise

    def _extract_scores(self, agent_output: str) -> Dict[str, float]:
        """Extract scores from agent output text"""
        scores = {}
        if "quality score" in agent_output:
            scores['content_quality'] = float(agent_output.split("quality score")[1].split(":")[1].split()[0])
        if "match score" in agent_output:
            scores['skills_match'] = float(agent_output.split("match score")[1].split(":")[1].split()[0])
        if "quality score" in agent_output:
            scores['experience_quality'] = float(agent_output.split("quality score")[1].split(":")[1].split()[0])
        return scores

    def _format_report_data(self, agent_outputs: str, total_usage: Dict[str, Any]) -> Dict[str, Any]:
        """Format agent outputs and usage stats for report template"""
        scores = self._extract_scores(agent_outputs)
        return {
            'content_quality_score': scores.get('content_quality', 0),
            'skills_match_score': scores.get('skills_match', 0),
            'experience_quality_score': scores.get('experience_quality', 0),
            'content_quality_analysis': agent_outputs.split("Content Quality Analysis:")[1].split("####")[0].strip(),
            'skills_analysis': agent_outputs.split("Skills Analysis:")[1].split("####")[0].strip(),
            'experience_analysis': agent_outputs.split("Experience Analysis:")[1].split("####")[0].strip(),
            'total_input_tokens': total_usage['total_input_tokens'],
            'total_output_tokens': total_usage['total_output_tokens'],
            'total_cost': round(total_usage['total_cost'], 2)
        }

    async def generate_report_pdf(self, agent_outputs: str, total_usage: Dict[str, Any]) -> str:
        """Generate a PDF report from agent outputs"""
        try:
            # Format the report data using the processor
            formatted_data = self.latex_processor.format_report_data(agent_outputs, total_usage)
            
            # Generate the PDF using the formatted data
            pdf_path = self.latex_processor.generate_report_pdf(
                template_name='report.tex.j2',
                data=formatted_data
            )
            logger.info(f"Successfully generated report PDF at: {pdf_path}")
            return pdf_path
        except Exception as e:
            logger.error(f"Error generating report PDF: {str(e)}")
            raise

    def _generate_docx(self, content: str, output_path: str) -> str:
        """Generate DOCX document using python-docx"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            # Create a new document
            doc = Document()
            
            # Set default styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Add content sections
            sections = content.split('#')
            for section in sections:
                if not section.strip():
                    continue
                    
                # Split section into title and content
                section_parts = section.split('\n', 1)
                title = section_parts[0].strip()
                body = section_parts[1].strip() if len(section_parts) > 1 else ''
                
                # Add section title
                heading = doc.add_heading(title, level=1)
                heading.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT (0)
                
                # Add section content
                if body:
                    for line in body.split('\n'):
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.style = doc.styles['Normal']
            
            # Save the document
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            raise RuntimeError(f"Failed to generate DOCX: {str(e)}")

    async def generate_resume_docx(self, resume_data: Dict[str, Any], personal_info: Dict[str, Any],
                                 job_title: str) -> Tuple[str, Dict[str, Any]]:
        """
        Generate a Word document resume using python-docx.
        Returns both the DOCX path and token usage statistics.
        """
        try:
            logger.info("Starting DOCX resume generation")
            
            # Use the same content formatting as PDF generation
            formatted_content = self.latex_processor.format_content(
                personal_info=personal_info,
                ai_content=resume_data['ai_content'],
                job_title=job_title,
                template_id=template_id
            )
            logger.info("Successfully formatted content using LaTeX processor")

            # Create output directory if it doesn't exist
            output_dir = Path(__file__).parent.parent / "output"
            output_dir.mkdir(exist_ok=True)
            
            # Generate filename
            timestamp = int(time.time())
            docx_path = str(output_dir / f"resume_{timestamp}.docx")
            
            # Generate DOCX using python-docx
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # Set smaller margins (0.5 inches on left/right)
            for section in doc.sections:
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Add name as title
            title = doc.add_heading(formatted_content.get('name', 'Resume'), 0)
            title.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER (1)
            
            # Add contact info
            contact_info = doc.add_paragraph()
            contact_info.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER (1)
            contact_info.add_run(
                f"{formatted_content.get('email', '')} | "
                f"{formatted_content.get('phone', '')} | "
                f"{formatted_content.get('location', '')}"
            )
            if formatted_content.get('linkedin'):
                contact_info.add_run(f" | {formatted_content['linkedin']}")
            
            # Add Professional Summary
            if formatted_content.get('summary'):
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(formatted_content['summary'])
            
            # Add Experience
            if formatted_content.get('experience'):
                doc.add_heading('Professional Experience', level=1)
                for exp in formatted_content['experience']:
                    if not exp:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{exp.get('title', '')} at {exp.get('company', '')}, {exp.get('duration', '')}").bold = True
                    if exp.get('achievements'):
                        for achievement in exp['achievements']:
                            if achievement:
                                doc.add_paragraph(achievement, style='List Bullet')
            
            # Add Projects
            if formatted_content.get('projects'):
                doc.add_heading('Projects', level=1)
                for project in formatted_content['projects']:
                    if not project:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(project.get('title', '')).bold = True
                    if project.get('highlights'):
                        for highlight in project['highlights']:
                            if highlight:
                                doc.add_paragraph(highlight, style='List Bullet')
            
            # Add Skills
            if formatted_content.get('skills'):
                doc.add_heading('Key Skills', level=1)
                skills_para = doc.add_paragraph()
                skills_para.add_run(' • '.join(formatted_content['skills']))
            
            # Add Education
            if formatted_content.get('education'):
                doc.add_heading('Education', level=1)
                for edu in formatted_content['education']:
                    if not edu:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{edu.get('degree', '')}\n").bold = True
                    p.add_run(f"{edu.get('institution', '')}\n")
                    p.add_run(f"{edu.get('year', '')}")
            
            # Add Certifications
            if formatted_content.get('certifications'):
                doc.add_heading('Certifications', level=1)
                for cert in formatted_content['certifications']:
                    if cert:
                        doc.add_paragraph(cert, style='List Bullet')
            
            # Add Achievements
            if formatted_content.get('achievements'):
                doc.add_heading('Achievements', level=1)
                for ach in formatted_content['achievements']:
                    if ach:
                        doc.add_paragraph(ach, style='List Bullet')
            
            # Save the document
            doc.save(docx_path)
            logger.info(f"Successfully generated DOCX at: {docx_path}")
            
            # Get the total token usage statistics
            total_usage = self.token_tracker.get_total_usage()
            
            return docx_path, total_usage

        except Exception as e:
            logger.error(f"Error generating DOCX resume: {str(e)}")
            raise

    def _format_docx_content(self, resume_data: Dict[str, Any], personal_info: Dict[str, Any],
                           job_title: str) -> str:
        """Format resume content for DOCX generation"""
        try:
            logger.info("Starting DOCX content formatting")
            logger.info(f"Resume data keys: {resume_data.keys()}")
            logger.info(f"Personal info keys: {personal_info.keys()}")
            
            if 'ai_content' not in resume_data:
                raise ValueError("ai_content not found in resume_data")
            
            content = resume_data['ai_content']
            logger.info("Processing content sections")
            
            # Split content into sections and process each section
            sections = {}
            current_section = None
            current_content = []
            
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                    
                if line.startswith('# '):
                    # If we were processing a section, save it
                    if current_section:
                        sections[current_section] = '\n'.join(current_content).strip()
                    # Start new section
                    current_section = line[2:].strip()
                    current_content = []
                elif line == '===':
                    continue  # Skip section markers
                else:
                    current_content.append(line)
            
            # Save the last section
            if current_section:
                sections[current_section] = '\n'.join(current_content).strip()
            
            logger.info(f"Found sections: {list(sections.keys())}")
            
            # Format the final content
            formatted_content = f"""
# {personal_info.get('name', 'Resume')}
{personal_info.get('email', '')} | {personal_info.get('phone', '')}
{personal_info.get('location', '')} | {personal_info.get('linkedin', '')}

# Professional Summary
{sections.get('Professional Summary', '')}

# Key Skills
{sections.get('Key Skills', '')}

# Professional Experience
{sections.get('Professional Experience', '')}

# Education
{sections.get('Education', '')}

# Projects
{sections.get('Projects', '')}

# Certifications & Achievements
{sections.get('Certifications & Achievements', '')}
"""
            logger.info("Successfully formatted DOCX content")
            return formatted_content.strip()
        except Exception as e:
            logger.error(f"Error formatting DOCX content: {str(e)}")
            raise RuntimeError(f"Failed to format DOCX content: {str(e)}")
        return content.strip()

    async def generate_resume(self, resume_data: Dict[str, Any], personal_info: Dict[str, Any],
                            job_title: str, format: str = 'pdf', template_id: str = None
                            ) -> Tuple[str, Dict[str, Any]]:
        """
        Generate a resume in the specified format (pdf or docx).
        
        Args:
            resume_data: Dictionary containing resume content
            personal_info: Dictionary containing personal information
            job_title: Target job title for the resume
            format: Output format ('pdf' or 'docx')
            template_id: ID of template to use (default: uses default template from config)
            
        Returns:
            Tuple of (file_path, token_usage_stats)
        """
        logger.info(f"Generating resume with template_id: {template_id}")
        if format.lower() == 'docx':
            return await self.generate_resume_docx(resume_data, personal_info, job_title)
        
        # Format the content for the LaTeX template
        try:
            formatted_content = self.latex_processor.format_content(
                personal_info=personal_info,
                ai_content=resume_data['ai_content'],
                job_title=job_title,
                template_id = template_id
            )

            # Generate PDF using LaTeX processor with selected template
            logger.info(f"Calling latex processor with template_id: {template_id}")
            pdf_result = self.latex_processor.generate_resume_pdf(
                content=formatted_content,
                template_id=template_id,
                user_id=str(personal_info["name"])  # Ensure user_id is string
            )
            
            # Handle new dict return format while maintaining backward compatibility
            pdf_path = pdf_result['pdf_path']
            if pdf_result.get('overflow', False):
                logger.warning(pdf_result['message'])
            
            logger.info(f"Generated PDF at {pdf_path} using template {template_id}")
            logger.info(f"Successfully generated PDF resume at: {pdf_path} using template: {template_id}")
            
            # Get the total token usage statistics
            total_usage = self.token_tracker.get_total_usage()
            
            return pdf_path, total_usage

        except Exception as e:
            logger.error(f"Error generating PDF resume: {str(e)}")
            raise
